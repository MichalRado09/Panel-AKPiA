import os
import tempfile
import time
import io
import json

import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- RDZEŃ DETERMINISTYCZNY (core/) ---
# Cały dobór i zliczanie dzieje się tutaj, NIE w LLM.
from core.ai_contract import build_extraction_prompt, parse_ai_json, build_response_schema
from core.parser import parse_devices, parse_ai_devices, devices_to_records, records_to_devices
from core.io_counter import count_io, format_balance, IO_TYPES
from core.plc_selector import select_plc, format_selection, PLATFORMY
from core.budget import calculate_budget, format_budget, GRUPY_RABATOWE
from core.cables import select_cables
from core.comparison import compare_variants
from core.scada_asix import select_asix
from core.cabinet import select_cabinet
from core.device_budget import build_device_budget, device_key, GRUPA_RABATOWA
from core.extraction_diff import compare_extractions, ExtractionDiff
from core.pdf_report import create_pdf_report
from core.validator import validate_offer, Severity
from core.hmi import build_hmi_selection, TYPOWE_PANELE
from core.signal_rules import NO_DATA

# --- 1. KONFIGURACJA BAZOWA ---
load_dotenv()

HISTORY_DIR = "historia_projektow"
OUTPUT_DIR = "outputs"
os.makedirs(HISTORY_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

GEMINI_MODEL = "gemini-3.5-flash"
MAX_INLINE_PDF_BYTES = 15 * 1024 * 1024

# Pliki "pamięci" aplikacji między sesjami — w .gitignore obok cennik.csv,
# bo mogą odzwierciedlać rzeczywiste ustawienia/rabaty/nazewnictwo firmy.
LAST_SETTINGS_FILE = "ustawienia_sesji.json"
LEARNED_SIGNALS_FILE = "nauczone_decyzje_sygnalow.json"


def load_last_settings() -> dict:
    """
    Ostatnio użyte ustawienia panelu bocznego (platforma/rezerwa/rabaty/...).
    Bez tego inżynier wpisywał te same, standardowe rabaty firmy od zera
    przy KAŻDYM uruchomieniu aplikacji — mimo że w praktyce zmieniają się
    rzadko. Brak pliku / uszkodzony JSON = po prostu wracamy do wbudowanych
    wartości domyślnych, nigdy nie wywalamy startu appki z tego powodu.
    """
    try:
        with open(LAST_SETTINGS_FILE, encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError, OSError):
        return {}


def save_last_settings(settings: dict) -> None:
    try:
        with open(LAST_SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
    except OSError:
        pass  # ustawienia to wygoda, nie krytyczna funkcja - błąd zapisu nie może wywalić UI


def load_learned_signal_decisions() -> dict:
    """
    Podpowiedzi dla sekcji "1b. Rozstrzygnij sygnały BRAK DANYCH", uczone
    z wcześniejszych decyzji inżyniera w TEJ instalacji aplikacji. Klucz to
    treść sygnału (np. "Pomiar temperatury - czujnik czy przetwornik?
    Rozstrzygnij ręcznie") — ta sama fraza wraca z core/device_rules.py
    w każdym projekcie, który trafi na ten sam niejednoznaczny wzorzec.

    WAŻNE - to WYŁĄCZNIE podpowiedź w UI (domyślny wybór w selectboxie),
    NIGDY automatyczna decyzja: inżynier zawsze musi kliknąć "Zastosuj",
    żeby sygnał faktycznie zmienił typ. Zero-hallucination zasada z
    core/*_rules.py ("nie zgadujemy") zostaje nienaruszona — to tylko
    szybsze podsunięcie tego, co inżynier sam wybrał poprzednio.
    """
    try:
        with open(LEARNED_SIGNALS_FILE, encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError, OSError):
        return {}


def save_learned_signal_decision(sygnal_nazwa: str, typ: str) -> None:
    decisions = load_learned_signal_decisions()
    decisions[sygnal_nazwa] = typ
    try:
        with open(LEARNED_SIGNALS_FILE, "w", encoding="utf-8") as f:
            json.dump(decisions, f, ensure_ascii=False, indent=2)
    except OSError:
        pass


def _get_secret(key: str) -> str | None:
    """
    Bezpieczny dostęp do st.secrets - zwraca None zamiast wywalać CAŁĄ
    aplikację, gdy nie istnieje ŻADEN plik secrets.toml.

    To jest realny, sprawdzony w praktyce przypadek: standardowy, udokumentowany
    sposób uruchomienia tej appki to sam plik .env, BEZ .streamlit/secrets.toml
    (ten drugi jest opisany jako alternatywa dla Streamlit Cloud). Nowsze wersje
    Streamlit (potwierdzone na 1.62.0) rzucają StreamlitSecretNotFoundError już
    na samym `"x" in st.secrets`, jeśli plik secrets.toml nie istnieje NIGDZIE -
    zamiast po cichu zwrócić False, jak można by się spodziewać po operatorze
    `in`. Bez tego zabezpieczenia get_api_key() wywala całą aplikację na starcie
    (błąd nieuchwycony przez check_password(), bo tam APP_PASSWORD z .env
    ratuje sytuację przez leniwe wyliczanie `and`).
    """
    try:
        if key in st.secrets:
            return st.secrets[key]
    except Exception:
        pass
    return None


# --- 2. ZABEZPIECZENIE APLIKACJI (LOGOWANIE) ---
def check_password():
    """Zwraca True, jeśli użytkownik wprowadził poprawne hasło."""
    correct_password = os.getenv("APP_PASSWORD") or _get_secret("APP_PASSWORD")

    if not correct_password:
        st.error("🚨 Krytyczny błąd: Brak skonfigurowanego hasła systemu! Ustaw APP_PASSWORD w .env lub secrets.toml.")
        st.stop()

    def password_entered():
        if st.session_state["password"] == correct_password:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.markdown("### 🔒 Panel Inżyniera AKPiA - Dostęp ograniczony")
        st.text_input("Wprowadź hasło dostępu:", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.markdown("### 🔒 Panel Inżyniera AKPiA - Dostęp ograniczony")
        st.text_input("Wprowadź hasło dostępu:", type="password", on_change=password_entered, key="password")
        st.error("😕 Niepoprawne hasło. Spróbuj ponownie.")
        return False

    return True


# --- 3. LOGIKA API GEMINI (tryb ekstrakcji JSON) ---
def normalize_model_name(model_name: str) -> str:
    model_name = model_name.strip()
    return model_name.removeprefix("models/") if model_name.startswith("models/") else model_name


def extract_response_text(response) -> str:
    if getattr(response, "text", None):
        return response.text
    if getattr(response, "candidates", None):
        candidate = response.candidates[0]
        content = getattr(candidate, "content", None)
        parts = getattr(content, "parts", None) if content else None
        if parts:
            text_parts = [part.text for part in parts if getattr(part, "text", None)]
            if text_parts:
                return "\n".join(text_parts)
    raise ValueError("Nie udało się odczytać treści odpowiedzi z API.")


def get_api_key() -> str:
    override = st.session_state.get("api_key_override", "").strip()
    if override:
        return override
    secret_key = _get_secret("GEMINI_API_KEY")
    if secret_key:
        return secret_key
    return os.getenv("GEMINI_API_KEY", "").strip()


def get_model_id() -> str:
    """
    Model Gemini do ekstrakcji - domyślnie stała GEMINI_MODEL, nadpisywalna
    na czas sesji w "Ustawienia API" (np. do przetestowania mocniejszego
    modelu bez edycji kodu i redeployu - dotąd jedyną drogą była zmiana
    stałej w kodzie, patrz README).
    """
    override = st.session_state.get("model_id_override", "").strip()
    return override or GEMINI_MODEL


def wait_for_file_active(client: genai.Client, file_name: str, timeout: int = 120) -> None:
    deadline = time.time() + timeout
    while time.time() < deadline:
        file_info = client.files.get(name=file_name)
        state_name = getattr(getattr(file_info, "state", None), "name", getattr(file_info, "state", None))
        if state_name == "ACTIVE":
            return
        if state_name == "FAILED":
            raise ValueError("Przetwarzanie pliku PDF w API zakończyło się błędem.")
        time.sleep(2)
    raise TimeoutError("Przekroczono czas oczekiwania na aktywację pliku w Gemini.")


def prepare_pdf_for_gemini(client: genai.Client, pdf_bytes: bytes):
    if len(pdf_bytes) <= MAX_INLINE_PDF_BYTES:
        pdf_part = types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf")
        return [pdf_part], None

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    try:
        uploaded_file = client.files.upload(file=tmp_path)
        wait_for_file_active(client, uploaded_file.name)
        return [uploaded_file], uploaded_file
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def call_gemini_with_retry(client, model_id, contents, system_instruction, response_schema=None, max_retries=3):
    for attempt in range(max_retries):
        try:
            config_kwargs = {"system_instruction": system_instruction}
            # Wymuszony tryb JSON: schema + mime_type = gwarancja poprawnego JSON
            if response_schema is not None:
                config_kwargs["response_mime_type"] = "application/json"
                config_kwargs["response_schema"] = response_schema
            response = client.models.generate_content(
                model=model_id,
                contents=contents,
                config=types.GenerateContentConfig(**config_kwargs),
            )
            return extract_response_text(response)
        except Exception as e:
            err_str = str(e)
            if ("503" in err_str or "429" in err_str) and attempt < max_retries - 1:
                time.sleep(5 * (attempt + 1))
                continue
            raise e


def run_extraction(api_key, excel_df, pdf_bytes, excel_filename, pdf_filename) -> list[dict]:
    """
    NOWY PRZEPŁYW: AI zwraca WYŁĄCZNIE listę urządzeń (JSON).
    Żadnego doboru ani BOM - to policzy rdzeń w core/.
    """
    client = genai.Client(api_key=api_key)
    model_id = normalize_model_name(get_model_id())
    sys_inst = build_extraction_prompt()

    user_prompt = "ZAŁĄCZONE ŹRÓDŁA DANYCH:\n"
    zrodla = []
    if excel_df is not None:
        zrodla.append(f"- Tabela urządzeń (Excel): {excel_filename}")
    if pdf_bytes is not None:
        zrodla.append(f"- Dokumentacja (PDF): {pdf_filename}")
    user_prompt += "\n".join(zrodla) + "\n\n"

    if excel_df is not None:
        user_prompt += (
            f"PODGLĄD DANYCH Z ARKUSZA EXCEL ({excel_filename}), format CSV:\n\n"
            f"{excel_df.to_csv(index=False)}\n\n"
        )

    contents = []
    uploaded_file = None
    if pdf_bytes is not None:
        pdf_parts, uploaded_file = prepare_pdf_for_gemini(client, pdf_bytes)
        contents.extend(pdf_parts)
        user_prompt += (
            f"Załączony PDF '{pdf_filename}' - wyodrębnij z niego urządzenia "
            "i połącz z danymi z Excela (bez podwójnego liczenia).\n\n"
        )

    user_prompt += "Zwróć listę urządzeń jako JSON zgodny ze schematem z instrukcji."
    contents.append(user_prompt)

    try:
        raw = call_gemini_with_retry(
            client, model_id, contents, sys_inst,
            response_schema=build_response_schema(),
        )
        return parse_ai_json(raw)
    finally:
        if uploaded_file:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass


# --- 4. GENEROWANIE PLIKÓW (na podstawie zatwierdzonych danych rdzenia) ---
def build_io_dataframe(devices) -> pd.DataFrame:
    """Tabela urządzeń z sygnałami - do podglądu i weryfikacji przez inżyniera."""
    rows = []
    for d in devices:
        sygnaly = "; ".join(f"{s['nazwa']} [{s['typ']}]" for s in d.sygnaly) or "-"
        zrodla = {s.get("source", "kolumna") for s in d.sygnaly}
        rows.append({
            "L.p.": d.lp,
            "Układ": d.uklad,
            "Oznaczenie": d.oznaczenie,
            "Opis": d.opis,
            "Ilość": d.ilosc,
            "Sygnały": sygnaly,
            "Źródło": ", ".join(sorted(zrodla)) if zrodla else "-",
            "Uwagi parsera": " | ".join(d.warnings) if d.warnings else "",
        })
    return pd.DataFrame(rows)


def create_word_report(devices, balance, project_label: str, platforma: str, rabaty: dict = None, hmi_entries: list = None, wycena_akpia_keys: set = None, price_overrides: dict = None) -> io.BytesIO:
    """Raport inżynierski z zatwierdzonym bilansem I/O (na razie: I/O; dobór w kolejnych modułach)."""
    doc = Document()
    title = doc.add_heading(f"Raport AKPiA: {project_label}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(f"Wygenerowano przez Panel Inżyniera AKPiA. Data: {time.strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading("Bilans sygnałów I/O", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Typ", "Baza", f"+Rezerwa {balance.reserve_percent}%"
    for t in IO_TYPES:
        cells = table.add_row().cells
        cells[0].text = t
        cells[1].text = str(balance.base[t])
        cells[2].text = str(balance.reserved[t])
    total = table.add_row().cells
    total[0].text, total[1].text, total[2].text = "RAZEM", str(balance.base_total), str(balance.reserved_total)

    doc.add_heading(f"Dobór sterownika ({platforma})", level=1)
    sel = select_plc(balance, platforma)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    h = tbl.rows[0].cells
    h[0].text, h[1].text, h[2].text = "Ilość", "Nr katalogowy", "Opis"
    for it in sel.items:
        c = tbl.add_row().cells
        c[0].text, c[1].text, c[2].text = str(it.ilosc), it.nr, it.opis

    hmi_sel_doc = build_hmi_selection(hmi_entries or [])
    if hmi_sel_doc.items:
        doc.add_heading("HMI — panele operatorskie lokalne", level=1)
        htbl = doc.add_table(rows=1, cols=3)
        htbl.style = "Light Grid Accent 1"
        hh = htbl.rows[0].cells
        hh[0].text, hh[1].text, hh[2].text = "Ilość", "Model", "Lokalizacja"
        for it in hmi_sel_doc.items:
            hc = htbl.add_row().cells
            hc[0].text, hc[1].text, hc[2].text = str(it.ilosc), it.nazwa, it.lokalizacja

    doc.add_heading("Kosztorys", level=1)
    budget = calculate_budget(sel.items, rabaty=rabaty or {})
    bt = doc.add_table(rows=1, cols=6)
    bt.style = "Light Grid Accent 1"
    bh = bt.rows[0].cells
    for i, h in enumerate(["Nr kat.", "Nazwa", "Ilość", "Kat. PLN", "Rabat %", "Netto PLN"]):
        bh[i].text = h
    for it in budget.items:
        bc = bt.add_row().cells
        bc[0].text = it.nr_katalogowy
        bc[1].text = it.nazwa
        bc[2].text = str(it.ilosc)
        bc[3].text = f"{it.cena_katalogowa:.2f}" if it.cena_katalogowa else "BRAK"
        bc[4].text = f"{it.rabat_pct:.0f}%"
        bc[5].text = f"{it.wartosc_netto:.2f}" if it.wartosc_netto else "-"
    sumr = bt.add_row().cells
    sumr[0].text = "SUMA"
    sumr[3].text = f"{budget.suma_katalogowa:.2f}"
    sumr[5].text = f"{budget.suma_netto:.2f}"
    if budget.brak_ceny:
        doc.add_paragraph(
            f"Uwaga: {len(budget.brak_ceny)} pozycji bez ceny katalogowej. "
            "Uzupełnij cennik, aby uzyskać pełny kosztorys."
        )

    dev_budget_doc = build_device_budget(devices, wycena_akpia_keys or set(), rabaty=rabaty or {},
                                          price_overrides=price_overrides or {})
    if dev_budget_doc.items:
        doc.add_heading("Kosztorys urządzeń AKPiA (wybór ręczny)", level=1)
        doc.add_paragraph(
            "Urządzenia obiektowe zaznaczone przez inżyniera jako wchodzące "
            "w zakres dostawy/wyceny AKPiA — osobno od sprzętu sterowniczego."
        )
        dt = doc.add_table(rows=1, cols=5)
        dt.style = "Light Grid Accent 1"
        dh = dt.rows[0].cells
        for i, h in enumerate(["Oznaczenie", "Opis", "Ilość", "Kat. PLN", "Netto PLN"]):
            dh[i].text = h
        for it in dev_budget_doc.items:
            dc = dt.add_row().cells
            dc[0].text = it.oznaczenie
            dc[1].text = it.opis
            dc[2].text = str(it.ilosc)
            dc[3].text = f"{it.cena_katalogowa:.2f}" if it.cena_katalogowa else "BRAK"
            dc[4].text = f"{it.wartosc_netto:.2f}" if it.wartosc_netto else "-"
        dsumr = dt.add_row().cells
        dsumr[0].text = "SUMA"
        dsumr[3].text = f"{dev_budget_doc.suma_katalogowa:.2f}"
        dsumr[4].text = f"{dev_budget_doc.suma_netto:.2f}"
        if dev_budget_doc.brak_ceny:
            doc.add_paragraph(
                f"Uwaga: {len(dev_budget_doc.brak_ceny)} pozycji bez ceny "
                "katalogowej (cennik nie zawiera jeszcze urządzeń obiektowych)."
            )

    doc.add_heading("Uwagi techniczne", level=1)
    if balance.undecided:
        doc.add_paragraph("Sygnały wymagające decyzji inżyniera (BRAK DANYCH):")
        for u in balance.undecided:
            doc.add_paragraph(f"• {u['urzadzenie']}: {u['sygnal']} (x{u['ilosc']})", style="List Bullet")
    inferred = balance.source_counts.get("typ_urzadzenia", 0)
    if inferred:
        doc.add_paragraph(
            f"Uwaga: {inferred} sygnał(ów) wywnioskowano z typu urządzenia "
            "(kolumny sygnałów były puste). Wymaga weryfikacji projektanta."
        )

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def create_devices_excel(devices, balance, platforma: str, rabaty: dict = None, cable_length: float = 25, asix_factor: float = 1.2, hmi_entries: list = None, wycena_akpia_keys: set = None, price_overrides: dict = None) -> io.BytesIO:
    """Excel: 9 zakładek — urządzenia, bilans, PLC, okablowanie, SCADA, porównanie, kosztorys, urządzenia AKPiA."""
    df_dev = build_io_dataframe(devices)
    df_io = pd.DataFrame(
        [{"Typ": t, "Baza": balance.base[t], f"Rezerwa {balance.reserve_percent}%": balance.reserved[t]}
         for t in IO_TYPES]
        + [{"Typ": "RAZEM", "Baza": balance.base_total, f"Rezerwa {balance.reserve_percent}%": balance.reserved_total}]
    )
    sel = select_plc(balance, platforma)
    df_plc = pd.DataFrame(
        [{"Ilość": it.ilosc, "Nr katalogowy": it.nr, "Opis": it.opis,
          "Grupa rabatowa": it.grupa_rabatowa}
         for it in sel.items]
    )
    cab = select_cables(devices, srednia_trasa_m=cable_length)
    df_cab = pd.DataFrame([
        {"Typ sygnału": it.typ_sygnalu, "Typ kabla": it.typ_kabla,
         "Urządzeń": it.ilosc_urzadzen, "Śr. trasa [m]": it.srednia_trasa_m,
         "Metraż [m]": it.metraz_m}
        for it in cab.items
    ] + [{"Typ sygnału": "RAZEM", "Metraż [m]": cab.total_metraz}])
    cab_sel = select_cabinet(balance, sel)
    df_cabinet = pd.DataFrame([
        {"Ilość": it.ilosc, "Nr katalogowy": it.nr_katalogowy, "Nazwa": it.nazwa,
         "Jednostka": it.jednostka, "Reguła doboru": it.uwaga,
         "Grupa rabatowa": it.grupa_rabatowa}
        for it in cab_sel.items
    ] + [
        {"Ilość": "", "Nr katalogowy": "BILANS PRĄDOWY 24V DC", "Nazwa": "", "Jednostka": "", "Reguła doboru": "", "Grupa rabatowa": ""},
        {"Ilość": cab_sel.prad_karty_ma, "Nr katalogowy": "Karty PLC [mA]", "Nazwa": "", "Jednostka": "mA", "Reguła doboru": "", "Grupa rabatowa": ""},
        {"Ilość": cab_sel.prad_przekazniki_ma, "Nr katalogowy": "Przekaźniki [mA]", "Nazwa": "", "Jednostka": "mA", "Reguła doboru": "", "Grupa rabatowa": ""},
        {"Ilość": cab_sel.prad_przetworniki_ma, "Nr katalogowy": "Przetworniki [mA]", "Nazwa": "", "Jednostka": "mA", "Reguła doboru": "", "Grupa rabatowa": ""},
        {"Ilość": cab_sel.prad_z_zapasem_a, "Nr katalogowy": "RAZEM z zapasem 30% [A]", "Nazwa": "", "Jednostka": "A", "Reguła doboru": "", "Grupa rabatowa": ""},
    ])
    asix = select_asix(balance, wspolczynnik=asix_factor)
    df_scada = pd.DataFrame([
        {"Parametr": "Sygnałów I/O (po rezerwie)", "Wartość": asix.zmienne_io},
        {"Parametr": f"Współczynnik zmiennych", "Wartość": asix.wspolczynnik},
        {"Parametr": "Zmiennych procesowych", "Wartość": asix.zmienne_obliczone},
        {"Parametr": "Pakiet licencyjny", "Wartość": asix.prog_nazwa},
        {"Parametr": "Sugestia architektury", "Wartość": asix.sugestia_opis},
    ] + [
        {"Parametr": f"Pozycja: {it.nr_katalogowy}", "Wartość": f"{it.ilosc}x {it.nazwa} ({it.cena_katalogowa} PLN)" if it.cena_katalogowa else f"{it.ilosc}x {it.nazwa}"}
        for it in asix.items
    ])
    variants = compare_variants(balance, rabaty=rabaty or {})
    df_cmp = pd.DataFrame([
        {"Platforma": v.platforma, "CPU": v.cpu,
         "Karty DI": v.karty_io.get("DI", 0), "Karty DO": v.karty_io.get("DO", 0),
         "Karty AI": v.karty_io.get("AI", 0), "Karty AO": v.karty_io.get("AO", 0),
         "Moduły łącznie": v.total_modules,
         "Suma netto [PLN]": v.suma_netto if v.suma_netto > 0 else None}
        for v in variants
    ])
    # Kosztorys: PLC + ASIX razem
    all_cost_items = list(sel.items)
    budget = calculate_budget(all_cost_items, rabaty=rabaty or {})
    # Dodaj pozycje ASIX i HMI do kosztorysu
    from core.plc_selector import PlcItem as _PI
    asix_plc_items = [_PI(nr=it.nr_katalogowy, opis=it.nazwa, ilosc=it.ilosc, grupa_rabatowa="ASIX") for it in asix.items]
    cab_plc_items = [_PI(nr=it.nr_katalogowy, opis=it.nazwa, ilosc=it.ilosc, grupa_rabatowa="APARATURA") for it in cab_sel.items]
    hmi_sel_xl = build_hmi_selection(hmi_entries or [])
    hmi_plc_items = [_PI(nr=f"HMI-{i}", opis=f"{it.nazwa} ({it.lokalizacja})" if it.lokalizacja else it.nazwa,
                         ilosc=it.ilosc, grupa_rabatowa="APARATURA")
                     for i, it in enumerate(hmi_sel_xl.items, 1)]
    df_hmi = pd.DataFrame([
        {"Ilość": it.ilosc, "Model": it.nazwa, "Lokalizacja": it.lokalizacja}
        for it in hmi_sel_xl.items
    ]) if hmi_sel_xl.items else pd.DataFrame([{"Ilość": "", "Model": "Brak dodanych paneli HMI", "Lokalizacja": ""}])
    budget_cab = calculate_budget(cab_plc_items, rabaty=rabaty or {})
    budget_asix = calculate_budget(asix_plc_items, rabaty=rabaty or {})
    budget_hmi = calculate_budget(hmi_plc_items, rabaty=rabaty or {})
    all_budget_items = budget.items + budget_asix.items + budget_cab.items + budget_hmi.items
    df_budget = pd.DataFrame([
        {"Nr katalogowy": it.nr_katalogowy, "Nazwa": it.nazwa, "Ilość": it.ilosc,
         "Jednostka": it.jednostka,
         "Cena katalogowa [PLN]": it.cena_katalogowa,
         "Rabat [%]": it.rabat_pct,
         "Cena netto/szt [PLN]": it.cena_netto_jed,
         "Wartość netto [PLN]": it.wartosc_netto}
        for it in all_budget_items
    ])
    dev_budget_xl = build_device_budget(devices, wycena_akpia_keys or set(), rabaty=rabaty or {},
                                         price_overrides=price_overrides or {})
    df_akpia_urz = pd.DataFrame([
        {"Oznaczenie": it.oznaczenie, "Opis": it.opis, "Ilość": it.ilosc,
         "Cena katalogowa [PLN]": it.cena_katalogowa, "Rabat [%]": it.rabat_pct,
         "Cena netto/szt [PLN]": it.cena_netto_jed, "Wartość netto [PLN]": it.wartosc_netto}
        for it in dev_budget_xl.items
    ] + [{"Oznaczenie": "SUMA", "Cena katalogowa [PLN]": dev_budget_xl.suma_katalogowa,
          "Wartość netto [PLN]": dev_budget_xl.suma_netto}]
    ) if dev_budget_xl.items else pd.DataFrame(
        [{"Oznaczenie": "", "Opis": "Brak zaznaczonych urządzeń (sekcja 1a w aplikacji)"}]
    )

    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df_dev.to_excel(writer, index=False, sheet_name="Urządzenia i sygnały")
        df_io.to_excel(writer, index=False, sheet_name="Bilans I-O")
        df_plc.to_excel(writer, index=False, sheet_name="Sterownik PLC")
        df_cab.to_excel(writer, index=False, sheet_name="Okablowanie")
        df_cabinet.to_excel(writer, index=False, sheet_name="Szafa SAKG")
        df_hmi.to_excel(writer, index=False, sheet_name="HMI")
        df_scada.to_excel(writer, index=False, sheet_name="SCADA ASIX")
        df_cmp.to_excel(writer, index=False, sheet_name="Porównanie wariantów")
        df_budget.to_excel(writer, index=False, sheet_name="Kosztorys")
        df_akpia_urz.to_excel(writer, index=False, sheet_name="Urządzenia AKPiA")
    bio.seek(0)
    return bio


def save_outputs_to_disk(
    project_label, devices, balance, word_bio, excel_bio,
    hmi_entries: list = None, wycena_akpia_keys: set = None, akpia_price_overrides: dict = None,
):
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    safe = project_label.replace("/", "_").replace("\\", "_")

    # Snapshot danych rdzenia (JSON) do historii - audytowalne, odtwarzalne.
    # HMI/wycena AKPiA/ceny ręczne trafiają do snapshotu, żeby "Wczytaj do
    # bieżącej analizy" (patrz Historia Projektów) mogło przywrócić PEŁNY
    # stan sesji, nie tylko listę urządzeń - inaczej Word/Excel zapisane
    # w tym momencie i sam snapshot rozjeżdżałyby się treścią.
    snapshot = {
        "project": project_label,
        "reserve_percent": balance.reserve_percent,
        "balance_base": balance.base,
        "balance_reserved": balance.reserved,
        "devices": devices_to_records(devices),
        "hmi_entries": hmi_entries or [],
        "wycena_akpia_keys": sorted(wycena_akpia_keys or set()),
        "akpia_price_overrides": akpia_price_overrides or {},
    }
    with open(os.path.join(HISTORY_DIR, f"{timestamp}_{safe}.json"), "w", encoding="utf-8") as f:
        json.dump(snapshot, f, ensure_ascii=False, indent=2)

    with open(os.path.join(OUTPUT_DIR, f"{timestamp}_{safe}.docx"), "wb") as f:
        f.write(word_bio.getvalue())
    with open(os.path.join(OUTPUT_DIR, f"{timestamp}_{safe}.xlsx"), "wb") as f:
        f.write(excel_bio.getvalue())


def persist_fresh_analysis(devices, project_label, settings) -> None:
    """
    Zapisuje snapshot do historii DOKŁADNIE RAZ, w momencie świeżej ekstrakcji
    (wywoływane wewnątrz handlera przycisku, więc uruchamia się tylko przy
    jego kliknięciu, nie przy każdym rerenderze Streamlit).

    Wcześniej zapis do historii wisiał na końcu main() poza handlerem
    przycisku, więc uruchamiał się przy KAŻDYM rerenderze strony - a
    Streamlit przelicza cały skrypt od nowa przy każdej interakcji (suwak
    rezerwy, checkbox wyceny AKPiA, dodanie pozycji HMI...). Efekt: jedna
    sesja pracy nad tym samym plikiem zaśmiecała historia_projektow/ i
    outputs/ dziesiątkami prawie identycznych snapshotów, myląc audyt,
    który ma pokazywać, co faktycznie zostało wysłane klientowi - nie
    każdy dotyk suwaka. Świadomy, kolejny zapis (np. po dopracowaniu
    rabatów) jest teraz przyciskiem "💾 Zapisz do historii" w sekcji 11.
    """
    balance = count_io(devices, reserve_percent=settings["reserve_percent"])
    word_bio = create_word_report(
        devices, balance, project_label, settings["platforma"], settings["rabaty"],
        st.session_state.get("hmi_entries", []), st.session_state.get("wycena_akpia_keys", set()),
        st.session_state.get("akpia_price_overrides", {}),
    )
    excel_bio = create_devices_excel(
        devices, balance, settings["platforma"], settings["rabaty"], settings["cable_length"],
        settings["asix_factor"], st.session_state.get("hmi_entries", []),
        st.session_state.get("wycena_akpia_keys", set()),
        st.session_state.get("akpia_price_overrides", {}),
    )
    save_outputs_to_disk(
        project_label, devices, balance, word_bio, excel_bio,
        hmi_entries=st.session_state.get("hmi_entries", []),
        wycena_akpia_keys=st.session_state.get("wycena_akpia_keys", set()),
        akpia_price_overrides=st.session_state.get("akpia_price_overrides", {}),
    )


# --- 5. INTERFEJS (STREAMLIT) ---
def list_excel_sheets(uploaded_file) -> list[str]:
    """Zwraca listę nazw arkuszy w pliku, bez wczytywania danych."""
    uploaded_file.seek(0)
    try:
        xl = pd.ExcelFile(uploaded_file)
        return xl.sheet_names
    except Exception:
        return []


def read_excel_file(uploaded_file, sheet_name=0) -> pd.DataFrame:
    """sheet_name: int (pozycja, domyślnie 0 = pierwszy) lub str (nazwa arkusza)."""
    try:
        uploaded_file.seek(0)
        return pd.read_excel(uploaded_file, sheet_name=sheet_name)
    except ImportError as exc:
        raise ValueError("Brak biblioteki do odczytu Excela (openpyxl/xlrd).") from exc
    except Exception as exc:
        raise ValueError(f"Nie udało się odczytać pliku Excel: {exc}") from exc


def build_project_label(excel_name, pdf_name, sheet_name=None) -> str:
    parts = [os.path.splitext(n)[0] for n in (excel_name, pdf_name) if n]
    label = "_".join(parts) if parts else "Projekt"
    # Dopisujemy nazwę arkusza, jeśli plik ma wiele arkuszy - zapobiega
    # sytuacji, w której dwa raporty z tego samego pliku (różne zakładki)
    # mają identyczną nazwę i łatwo je pomylić.
    if isinstance(sheet_name, str):
        safe_sheet = sheet_name.replace("/", "-").replace("\\", "-")
        label = f"{label}_{safe_sheet}"
    return label


def _clamp(value, lo, hi, default):
    try:
        value = type(default)(value)
    except (TypeError, ValueError):
        return default
    return max(lo, min(hi, value))


def render_sidebar():
    # Ostatnio użyte ustawienia (patrz load_last_settings) - żeby inżynier
    # nie wpisywał tych samych, standardowych rabatów firmy od zera przy
    # każdym uruchomieniu appki. Wartości clampowane na wypadek ręcznie
    # popsutego pliku - nigdy nie mają wywalić startu UI.
    last = load_last_settings()

    st.sidebar.title("🛠 Panel Inżyniera AKPiA")
    st.sidebar.markdown("---")
    page = st.sidebar.radio("Tryb działania",
                            ["Analiza Projektu", "Historia Projektów", "Ustawienia API"],
                            key="tryb_dzialania")
    st.sidebar.markdown("---")
    st.sidebar.markdown("**Parametry Techniczne**")
    platformy_lista = list(PLATFORMY.keys())
    domyslna_platforma = last.get("platforma")
    platforma = st.sidebar.selectbox(
        "Platforma PLC", platformy_lista,
        index=platformy_lista.index(domyslna_platforma) if domyslna_platforma in platformy_lista else 0,
    )
    reserve_percent = st.sidebar.slider(
        "Rezerwa sprzętowa [%]", 0, 100, _clamp(last.get("reserve_percent"), 0, 100, 30), 5)
    cable_length = st.sidebar.slider(
        "Średnia trasa kablowa [m]", 5, 200, _clamp(last.get("cable_length"), 5, 200, 25), 5,
        help="Jedna średnia dla wszystkich typów sygnału. Zmierzone na realnej "
             "liście kablowej DPK2 Wujek (196 kabli): analogi ~65 m, DI ~36 m, "
             "DO ~34 m, falowniki ~46 m, ethernet ~17 m. Trasy różnią się po "
             "typie prawie dwukrotnie — najdłuższe są analogi, bo przetworniki "
             "stoją w terenie. Ustaw wartość bliską temu, co dominuje w Twoim "
             "projekcie; 25 m odpowiada właściwie tylko sieci w szafie.")
    asix_factor = st.sidebar.slider(
        "Współczynnik zmiennych ASIX", 1.0, 2.0, _clamp(last.get("asix_factor"), 1.0, 2.0, 1.2), 0.1)

    st.sidebar.markdown("---")
    st.sidebar.markdown("**Rabaty firmowe [%]**")
    rabaty = {}
    last_rabaty = last.get("rabaty") or {}
    for grupa, default in GRUPY_RABATOWE.items():
        rabaty[grupa] = st.sidebar.number_input(
            f"{grupa}", min_value=0, max_value=100,
            value=_clamp(last_rabaty.get(grupa), 0, 100, default), step=1, key=f"rabat_{grupa}"
        )

    st.sidebar.markdown("---")
    if st.sidebar.button("Wyloguj", width="stretch"):
        st.session_state.clear()
        st.rerun()

    settings = {"reserve_percent": reserve_percent, "platforma": platforma, "rabaty": rabaty,
                "cable_length": cable_length, "asix_factor": asix_factor}
    save_last_settings(settings)
    return page, settings


def render_extraction_diff_panel() -> None:
    """
    Panel porównania ścieżki OFFLINE i AI - widoczny po użyciu przycisku
    '⚖ Policz + zweryfikuj przez AI'. Pokazuje deterministyczne różnice
    (core.extraction_diff) jako pomoc diagnostyczną, NIGDY jako podstawę
    do automatycznej zmiany liczb. Wybór, którą wersję zatwierdzić do
    dalszej pracy, zawsze należy do inżyniera (przyciski niżej).
    """
    diff: ExtractionDiff = st.session_state.extraction_diff

    st.subheader("⚖ Porównanie: OFFLINE vs AI")

    if diff.identyczne:
        st.success(
            f"✓ Obie ścieżki dają identyczny wynik: {diff.liczba_urzadzen_offline} "
            "urządzeń, identyczny bilans I/O. Ekstrakcja AI zgodna z parserem offline."
        )
        return

    st.warning(
        f"⚠ Wykryto różnice — offline: {diff.liczba_urzadzen_offline} urządzeń, "
        f"AI: {diff.liczba_urzadzen_ai} urządzeń."
    )

    delta_cols = st.columns(4)
    for i, t in enumerate(IO_TYPES):
        d = diff.balans_delta[t]
        delta_cols[i].metric(t, f"{d:+d}" if d != 0 else "0",
                             help="Delta AI minus offline")

    def _render_lista(tytul: str, wpisy: list) -> None:
        if not wpisy:
            return
        st.markdown(f"**{tytul} ({len(wpisy)}):**")
        for e in wpisy:
            st.markdown(f"- **{e.oznaczenie}**: {e.opis} (x{e.ilosc})")

    _render_lista("Tylko w OFFLINE", diff.tylko_w_offline)
    _render_lista("Tylko w AI", diff.tylko_w_ai)

    st.caption(
        "Decyzję, którą wersję zatwierdzić do dalszej pracy, podejmuje inżynier. "
        "Domyślnie do sekcji poniżej trafia wynik OFFLINE (deterministyczny)."
    )
    if st.session_state.get("devices_ai_alternative") is not None:
        if st.button("🔄 Użyj zamiast tego wyniku AI", width="content"):
            devices_ai_wybrane = st.session_state.devices_ai_alternative
            st.session_state.devices = devices_ai_wybrane
            st.rerun()


def render_device_table_editor(devices: list) -> list:
    """
    Edytowalna tabela urządzeń: inżynier może poprawić błędnie odczytaną
    ILOŚĆ oraz USUNĄĆ pozycję (np. duplikat, błędnie wyekstrahowany wiersz)
    wprost w interfejsie — bez wracania do źródłowego pliku i ponownego
    przechodzenia całym przepływem (upload -> parsowanie -> ekstrakcja...).
    To był dotąd realny brak: tabela w sekcji 1 była WYŁĄCZNIE do odczytu
    (st.dataframe), więc każda drobna pomyłka parsera/AI wymagała edycji
    Excela i powtórzenia analizy od zera.

    Świadomie NIE pozwala edytować Opisu/Oznaczenia/Układu w tej tabeli —
    zmiana opisu nie przeliczyłaby ponownie reguły typu urządzenia (sygnały
    są przypisywane raz, przy parsowaniu w core/parser.py), więc cicha
    edycja opisu bez przeliczenia sygnałów tworzyłaby niespójność, którą
    trudno zauważyć. Klasyfikację sygnałów BRAK DANYCH rozstrzyga osobna
    sekcja niżej (render_undecided_signal_resolver).

    Zwraca zaktualizowaną listę Device — wywołujący ma podmienić nią
    st.session_state.devices, żeby usunięcie wiersza przetrwało rerender.
    """
    rows = []
    for i, d in enumerate(devices):
        sygnaly = "; ".join(f"{s['nazwa']} [{s['typ']}]" for s in d.sygnaly) or "-"
        zrodla = {s.get("source", "kolumna") for s in d.sygnaly}
        rows.append({
            "L.p.": d.lp,
            "Układ": d.uklad,
            "Oznaczenie": d.oznaczenie,
            "Opis": d.opis,
            "Ilość": d.ilosc,
            "Sygnały": sygnaly,
            "Źródło": ", ".join(sorted(zrodla)) if zrodla else "-",
            "Uwagi parsera": " | ".join(d.warnings) if d.warnings else "",
            "_key": device_key(d, i),
        })
    df = pd.DataFrame(rows)

    edited = st.data_editor(
        df,
        column_config={
            "Ilość": st.column_config.NumberColumn(
                "Ilość", min_value=0, step=1,
                help="Popraw, jeśli parser/AI źle odczytał liczbę sztuk.",
            ),
            "_key": None,  # ukrywa kolumnę techniczną w UI
        },
        disabled=["L.p.", "Układ", "Oznaczenie", "Opis", "Sygnały", "Źródło", "Uwagi parsera"],
        num_rows="dynamic",
        hide_index=True,
        width="stretch",
        key="device_table_editor",
    )

    by_key = {device_key(d, i): d for i, d in enumerate(devices)}
    updated: list = []
    n_new_ignored = 0
    for _, row in edited.iterrows():
        dev = by_key.get(row["_key"])
        if dev is None:
            # Ręcznie dodany, pusty wiersz w edytorze — nie da się dla niego
            # zbudować sygnałów I/O bez opisu, więc pomijamy go zamiast
            # cicho liczyć "urządzenie" bez żadnych danych.
            n_new_ignored += 1
            continue
        try:
            dev.ilosc = max(0, int(row["Ilość"]))
        except (TypeError, ValueError):
            pass
        updated.append(dev)

    if n_new_ignored:
        st.warning(
            f"Zignorowano {n_new_ignored} ręcznie dodany wiersz — dodawanie nowych "
            "urządzeń w tej tabeli nie jest wspierane (brak opisu = brak reguły "
            "sygnałów). Dodaj urządzenie w źródłowym pliku i wczytaj ponownie."
        )
    if len(updated) != len(devices):
        st.caption(
            f"ℹ Usunięto {len(devices) - len(updated)} pozycję/e z listy — "
            "zmiana obowiązuje do końca tej analizy (nie zmienia pliku źródłowego)."
        )

    return updated


def render_undecided_signal_resolver(devices: list) -> None:
    """
    Ręczne rozstrzygnięcie sygnałów BRAK DANYCH (nierozpoznany typ DI/DO/AI/AO)
    WPROST w interfejsie — bez edycji źródłowego pliku. Dotąd jedynym sposobem
    na naprawienie takiego sygnału było wyjście z aplikacji, poprawienie opisu
    w Excelu i ponowne przejście całej analizy — mimo że dane wymagające
    decyzji są dokładnie znane (patrz core/signal_rules.py, core/device_rules.py:
    "NIE zgadujemy").

    Decyzja nadpisuje typ TYLKO tego jednego sygnału (mutacja in-place na
    obiekcie Device przechowywanym w st.session_state.devices) — reszta
    danych urządzenia zostaje bez zmian, bilans I/O przelicza się przy
    najbliższym rerenderze. Decyzja jest też zapamiętywana (patrz
    save_learned_signal_decision) jako PODPOWIEDŹ do następnego razu, gdy
    ten sam sygnał („Pomiar temperatury - czujnik czy przetwornik?...")
    pojawi się w innym projekcie — inżynier wciąż musi kliknąć "Zastosuj",
    podpowiedź tylko wstępnie zaznacza wybór.
    """
    undecided = [
        (i, si, dev, sig)
        for i, dev in enumerate(devices)
        for si, sig in enumerate(dev.sygnaly)
        if sig.get("typ") == NO_DATA
    ]
    if not undecided:
        return

    learned = load_learned_signal_decisions()
    typy_opcje = ["— nie rozstrzygnięto —", "DI", "DO", "AI", "AO"]

    st.subheader("1b. Rozstrzygnij sygnały bez klasyfikacji (BRAK DANYCH)")
    st.caption(
        f"{len(undecided)} sygnał(ów) nie ma jednoznacznej klasyfikacji DI/DO/AI/AO "
        "(np. „czujnik czy przetwornik?”) i nie wchodzi do bilansu I/O, dopóki nie "
        "zostaną rozstrzygnięte. Rozstrzygnij poniżej, żeby nie edytować pliku "
        "źródłowego tylko dla tej jednej decyzji. Podpowiedź (jeśli jest) pochodzi "
        "z Twojej wcześniejszej decyzji dla tego samego sygnału w innym projekcie — "
        "i tak wymaga kliknięcia „Zastosuj”."
    )
    for i, si, dev, sig in undecided:
        cols = st.columns([3, 2, 1])
        nazwa_sygnalu = sig.get("nazwa", "")
        cols[0].markdown(f"**{dev.oznaczenie or dev.opis}** — {nazwa_sygnalu}")
        podpowiedz = learned.get(nazwa_sygnalu)
        wybor = cols[1].selectbox(
            "Typ sygnału", typy_opcje,
            index=typy_opcje.index(podpowiedz) if podpowiedz in typy_opcje else 0,
            key=f"undecided_{i}_{si}", label_visibility="collapsed",
        )
        if cols[2].button("Zastosuj", key=f"undecided_apply_{i}_{si}", width="stretch"):
            if wybor != "— nie rozstrzygnięto —":
                dev.sygnaly[si]["typ"] = wybor
                dev.sygnaly[si]["source"] = "inzynier"
                save_learned_signal_decision(nazwa_sygnalu, wybor)
                st.rerun()


def render_device_budget_selector(devices, rabaty: dict) -> None:
    """
    Checkbox-lista urządzeń obiektowych do RĘCZNEGO oznaczenia, które wchodzą
    w zakres wyceny AKPiA (typowo: przetworniki pomiarowe). Stan trzymany w
    st.session_state, kluczowany przez device_key() - przetrwa przeliczenia
    w obrębie tej samej sesji, dopóki lista urządzeń się nie zmieni.

    Kolumna "Cena ręczna [PLN]" pozwala wpisać cenę wprost tutaj, gdy
    cennik.csv nie ma jeszcze dopasowania po oznaczeniu/opisie (typowy stan
    dla urządzeń obiektowych - patrz core/device_budget.py) - bez tego
    jedynym sposobem na wycenę takiej pozycji było ręczne dopisanie wiersza
    do cennik.csv poza aplikacją.
    """
    if "wycena_akpia_keys" not in st.session_state:
        st.session_state.wycena_akpia_keys = set()
    if "akpia_price_overrides" not in st.session_state:
        st.session_state.akpia_price_overrides = {}

    rows = []
    for i, d in enumerate(devices):
        key = device_key(d, i)
        rows.append({
            "Wycena AKPiA": key in st.session_state.wycena_akpia_keys,
            "Oznaczenie": d.oznaczenie or "-",
            "Opis": d.opis,
            "Ilość": d.ilosc,
            "Cena ręczna [PLN]": st.session_state.akpia_price_overrides.get(key),
            "_key": key,  # ukryta kolumna pomocnicza, nie do edycji
        })
    df_sel = pd.DataFrame(rows)

    edited = st.data_editor(
        df_sel,
        column_config={
            "Wycena AKPiA": st.column_config.CheckboxColumn(
                "Wycena AKPiA", help="Zaznacz, jeśli to urządzenie ma trafić do kosztorysu AKPiA"
            ),
            "Cena ręczna [PLN]": st.column_config.NumberColumn(
                "Cena ręczna [PLN]", min_value=0.0, step=1.0,
                help="Cena katalogowa za sztukę, jeśli cennik.csv nie ma dopasowania. "
                     "Puste = szukaj w cenniku (BRAK CENY, jeśli i tam nic nie ma).",
            ),
            "_key": None,  # ukrywa kolumnę techniczną w UI
        },
        disabled=["Oznaczenie", "Opis", "Ilość"],
        hide_index=True,
        width="stretch",
        key="device_budget_editor",
    )

    # Synchronizacja stanu na podstawie tego, co inżynier zaznaczył/wpisał w tabeli
    st.session_state.wycena_akpia_keys = set(
        edited.loc[edited["Wycena AKPiA"], "_key"]
    )
    st.session_state.akpia_price_overrides = {
        row["_key"]: float(row["Cena ręczna [PLN]"])
        for _, row in edited.iterrows()
        if pd.notna(row["Cena ręczna [PLN]"])
    }

    dev_budget = build_device_budget(
        devices, st.session_state.wycena_akpia_keys, rabaty=rabaty,
        price_overrides=st.session_state.akpia_price_overrides,
    )
    if dev_budget.items:
        st.caption(f"Zaznaczono {len(dev_budget.items)} pozycji do kosztorysu AKPiA.")
    else:
        st.caption("Brak zaznaczonych pozycji — żadne urządzenie obiektowe nie trafi do kosztorysu.")


def render_results(devices, balance, project_label, platforma, rabaty, cable_length=25, asix_factor=1.2):
    """Sekcja HITL: urządzenia + bilans + PLC + okablowanie + SCADA + porównanie + kosztorys + pliki."""
    st.subheader("1. Zidentyfikowane urządzenia (do weryfikacji)")
    st.caption("Kolumna 'Źródło' pokazuje, czy sygnał pochodzi z danych (kolumna), "
               "czy z reguły typu urządzenia. Popraw Ilość albo usuń błędną pozycję "
               "wprost w tabeli — zmiana obowiązuje od razu, bez ponownego wgrywania pliku.")
    updated_devices = render_device_table_editor(devices)
    if len(updated_devices) != len(devices):
        st.session_state.devices = updated_devices
        st.rerun()
    devices = updated_devices

    render_undecided_signal_resolver(devices)

    st.subheader("1a. Urządzenia obiektowe wchodzące w zakres wyceny AKPiA")
    st.caption(
        "Większość urządzeń obiektowych (pompy, zawory, siłowniki) fizycznie "
        "znajduje się na hali i jest dostarczana/wyceniana poza automatyką - "
        "program ich NIE wycenia automatycznie. Zaznacz ręcznie te pozycje "
        "(typowo: przetworniki pomiarowe), które WCHODZĄ w zakres dostawy "
        "AKPiA i mają trafić do kosztorysu."
    )
    render_device_budget_selector(devices, rabaty)

    st.subheader("2. Bilans sygnałów I/O")
    cols = st.columns(len(IO_TYPES) + 1)
    for i, t in enumerate(IO_TYPES):
        cols[i].metric(t, balance.base[t], f"+{balance.reserved[t] - balance.base[t]} rez.")
    cols[-1].metric("RAZEM", balance.base_total, f"→ {balance.reserved_total}")

    if balance.undecided:
        st.warning(f"⚠ {len(balance.undecided)} sygnał(ów) BRAK DANYCH - wymaga decyzji inżyniera "
                   f"(szczegóły w raporcie).")
    inferred = balance.source_counts.get("typ_urzadzenia", 0)
    if inferred:
        st.info(f"ℹ {inferred} sygnał(ów) wywnioskowano z typu urządzenia (puste kolumny). Zweryfikuj.")

    st.subheader(f"3. Dobór sterownika ({platforma})")
    st.caption("Reguły zweryfikowane na realnych projektach (Wujek, TOM, Malbork).")
    try:
        sel = select_plc(balance, platforma)
    except FileNotFoundError:
        # Jedyny plik danych bez łagodnego fallbacku (patrz WYMAGANE_PLIKI.md) -
        # bez tego inżynier dostawałby surowy traceback Streamlit zamiast
        # zrozumiałej informacji, co dokładnie brakuje i jak to naprawić.
        st.error(
            f"🚨 Brak pliku katalogu kart dla platformy „{platforma}” w katalogu "
            f"`katalogi/`. Sprawdź core/plc_selector.py::PLATFORMY i czy "
            f"odpowiadający plik CSV faktycznie tam jest — patrz WYMAGANE_PLIKI.md."
        )
        st.stop()
    df_plc = pd.DataFrame([
        {"Ilość": it.ilosc, "Nr katalogowy": it.nr, "Opis": it.opis} for it in sel.items
    ])
    st.dataframe(df_plc, width="stretch")
    util_cols = st.columns(len(IO_TYPES))
    for i, t in enumerate(IO_TYPES):
        u = sel.utilization.get(t, {})
        util_cols[i].metric(f"{t} karty", u.get("kart", 0),
                            f"zapas {u.get('zapas_kanałów', 0)} kan.")

    st.subheader("4. Zestawienie kablowe")
    cab = select_cables(devices, srednia_trasa_m=cable_length)
    df_cab = pd.DataFrame([
        {"Typ sygnału": it.typ_sygnalu, "Typ kabla": it.typ_kabla,
         "Urządzeń": it.ilosc_urzadzen, "Metraż [m]": it.metraz_m}
        for it in cab.items
    ])
    st.dataframe(df_cab, width="stretch")
    st.metric("Razem metraż", f"{cab.total_metraz:.0f} m",
              f"(śr. trasa {cable_length}m, naddatek +15%)")

    st.subheader("5. SCADA ASIX")
    asix = select_asix(balance, wspolczynnik=asix_factor)
    scada_cols = st.columns(3)
    scada_cols[0].metric("Sygnałów I/O", asix.zmienne_io)
    scada_cols[1].metric("Zmiennych procesowych", asix.zmienne_obliczone,
                         f"×{asix.wspolczynnik}")
    scada_cols[2].metric("Pakiet licencyjny", asix.prog_nazwa)
    st.info(f"💡 Sugestia architektury: {asix.sugestia_opis}")
    if asix.items:
        df_asix = pd.DataFrame([
            {"Nr katalogowy": it.nr_katalogowy, "Nazwa": it.nazwa,
             "Ilość": it.ilosc,
             "Cena kat. [PLN]": f"{it.cena_katalogowa:.2f}" if it.cena_katalogowa else "BRAK"}
            for it in asix.items
        ])
        st.dataframe(df_asix, width="stretch")
    for w in asix.warnings:
        st.warning(w)

    st.subheader("6. HMI — panele operatorskie lokalne")
    st.caption("Traktowane OSOBNO od SCADA/ASIX (wytyczna: HMI = sterowanie lokalne, "
               "ASIX = stacja operatorska/serwer). Wybór manualny — brak zwalidowanego "
               "wzorca doboru w dokumentacji firmy, więc inżynier dodaje pozycje ręcznie.")

    if "hmi_entries" not in st.session_state:
        st.session_state.hmi_entries = []

    with st.form("hmi_add_form", clear_on_submit=True):
        hc1, hc2, hc3, hc4 = st.columns([3, 1, 2, 1])
        with hc1:
            hmi_nazwa = st.selectbox("Model panelu", TYPOWE_PANELE, key="hmi_model_select")
            hmi_nazwa_custom = st.text_input("...lub wpisz własny model", key="hmi_model_custom")
        with hc2:
            hmi_ilosc = st.number_input("Ilość", min_value=1, value=1, key="hmi_ilosc_input")
        with hc3:
            hmi_lok = st.text_input("Lokalizacja", key="hmi_lok_input",
                                    placeholder="np. szafa sterownicza")
        with hc4:
            st.markdown("&nbsp;")
            hmi_submit = st.form_submit_button("➕ Dodaj")
        if hmi_submit:
            final_nazwa = hmi_nazwa_custom.strip() or hmi_nazwa
            if final_nazwa and final_nazwa != "Inny / wpisz ręcznie":
                st.session_state.hmi_entries.append(
                    {"nazwa": final_nazwa, "ilosc": hmi_ilosc, "lokalizacja": hmi_lok}
                )

    hmi_sel = build_hmi_selection(st.session_state.hmi_entries)
    if hmi_sel.items:
        df_hmi = pd.DataFrame([
            {"Ilość": it.ilosc, "Model": it.nazwa, "Lokalizacja": it.lokalizacja}
            for it in hmi_sel.items
        ])
        st.dataframe(df_hmi, width="stretch")
        if st.button("🗑 Wyczyść listę HMI"):
            st.session_state.hmi_entries = []
            st.rerun()
    else:
        st.caption("Brak dodanych paneli HMI — poprawny stan, jeśli projekt ich nie wymaga.")

    st.subheader("7. Szafa sterownicza (+SAKG)")
    cab_sel = select_cabinet(balance, sel)
    df_cab_items = pd.DataFrame([
        {"Ilość": it.ilosc, "Nr katalogowy": it.nr_katalogowy,
         "Nazwa": it.nazwa, "Reguła": it.uwaga}
        for it in cab_sel.items
    ])
    st.dataframe(df_cab_items, width="stretch")
    pw = st.columns(4)
    pw[0].metric("Karty PLC", f"{cab_sel.prad_karty_ma} mA")
    pw[1].metric("Przekaźniki", f"{cab_sel.prad_przekazniki_ma} mA")
    pw[2].metric("Przetworniki", f"{cab_sel.prad_przetworniki_ma} mA")
    pw[3].metric("Zasilacz 24V", f"{cab_sel.zasilacz_a} A",
                 f"bilans {cab_sel.prad_z_zapasem_a} A")
    for w in cab_sel.warnings:
        st.info(f"ℹ {w}")

    st.subheader("8. Porównanie wariantów sterowników")
    st.caption("Ten sam bilans I/O — trzy platformy obok siebie.")
    variants = compare_variants(balance, rabaty=rabaty)
    df_cmp = pd.DataFrame([
        {"Platforma": v.platforma, "CPU": v.cpu,
         "Karty DI": v.karty_io.get("DI", 0), "Karty DO": v.karty_io.get("DO", 0),
         "Karty AI": v.karty_io.get("AI", 0), "Karty AO": v.karty_io.get("AO", 0),
         "Moduły łącznie": v.total_modules,
         "Netto [PLN]": f"{v.suma_netto:,.2f}" if v.suma_netto > 0 else "brak cen"}
        for v in variants
    ])
    st.dataframe(df_cmp, width="stretch")

    st.subheader("9. Kosztorys")
    budget = calculate_budget(sel.items, rabaty=rabaty)
    df_budget = pd.DataFrame([
        {
            "Nr katalogowy": it.nr_katalogowy,
            "Nazwa": it.nazwa,
            "Ilość": it.ilosc,
            "Cena kat. [PLN]": f"{it.cena_katalogowa:.2f}" if it.cena_katalogowa else "BRAK",
            "Rabat [%]": f"{it.rabat_pct:.0f}",
            "Netto/szt [PLN]": f"{it.cena_netto_jed:.2f}" if it.cena_netto_jed else "-",
            "Wartość netto [PLN]": f"{it.wartosc_netto:.2f}" if it.wartosc_netto else "-",
        }
        for it in budget.items
    ])
    st.dataframe(df_budget, width="stretch")

    sum_cols = st.columns(2)
    sum_cols[0].metric("Suma katalogowa", f"{budget.suma_katalogowa:,.2f} PLN")
    sum_cols[1].metric("Suma netto (po rabatach)", f"{budget.suma_netto:,.2f} PLN")

    if budget.brak_ceny:
        st.warning(f"⚠ {len(budget.brak_ceny)} pozycji bez ceny katalogowej — "
                   "uzupełnij cennik, aby uzyskać pełny kosztorys.")

    st.subheader("9a. Kosztorys urządzeń AKPiA (wybór ręczny)")
    st.caption("Pozycje zaznaczone w sekcji 1a — osobno od sprzętu sterowniczego, "
               "bo dotyczą urządzeń obiektowych (np. przetworników), a nie kart PLC/szafy/SCADA.")
    dev_budget = build_device_budget(
        devices, st.session_state.get("wycena_akpia_keys", set()), rabaty=rabaty,
        price_overrides=st.session_state.get("akpia_price_overrides", {}),
    )
    if dev_budget.items:
        df_dev_budget = pd.DataFrame([
            {
                "Oznaczenie": it.oznaczenie,
                "Opis": it.opis,
                "Ilość": it.ilosc,
                "Cena kat. [PLN]": f"{it.cena_katalogowa:.2f}" if it.cena_katalogowa else "BRAK",
                "Rabat [%]": f"{it.rabat_pct:.0f}",
                "Wartość netto [PLN]": f"{it.wartosc_netto:.2f}" if it.wartosc_netto else "-",
            }
            for it in dev_budget.items
        ])
        st.dataframe(df_dev_budget, width="stretch")
        sum_cols2 = st.columns(2)
        sum_cols2[0].metric("Suma katalogowa (AKPiA)", f"{dev_budget.suma_katalogowa:,.2f} PLN")
        sum_cols2[1].metric("Suma netto (AKPiA)", f"{dev_budget.suma_netto:,.2f} PLN")
        if dev_budget.brak_ceny:
            st.warning(f"⚠ {len(dev_budget.brak_ceny)} pozycji bez ceny katalogowej — "
                       "cennik nie zawiera jeszcze urządzeń obiektowych, uzupełnij ręcznie.")
    else:
        st.caption("Brak zaznaczonych urządzeń — sekcja 1a pozwala je dodać.")

    st.subheader("10. Weryfikacja kompletności oferty")
    # Reużywamy sel/cab/cab_sel/asix/budget policzone wyżej (sekcje 3-9) —
    # bez tego walidator liczył PLC/kable/SCADA/kosztorys jeszcze raz, mimo
    # że wynik jest identyczny (te same wejścia), tylko po to, żeby za chwilę
    # go wyrzucić.
    val_report = validate_offer(devices, balance, sel, cab, cab_sel, asix, budget)

    if val_report.is_clean:
        st.success("✓ Brak zastrzeżeń — oferta wygląda na spójną.")
    else:
        for issue in val_report.errors:
            st.error(f"🔴 **[{issue.category}]** {issue.message}")
        for issue in val_report.warnings:
            st.warning(f"🟡 **[{issue.category}]** {issue.message}")
        for issue in val_report.infos:
            st.info(f"ℹ️ **[{issue.category}]** {issue.message}")
        st.caption("Powyższe to sugestie do weryfikacji — nie blokują generowania dokumentów. "
                   "Ostateczna decyzja należy do inżyniera.")

    st.subheader("11. Pobierz dokumenty")
    word_bio = create_word_report(devices, balance, project_label, platforma, rabaty, st.session_state.get("hmi_entries", []), st.session_state.get("wycena_akpia_keys", set()))
    excel_bio = create_devices_excel(devices, balance, platforma, rabaty, cable_length, asix_factor, st.session_state.get("hmi_entries", []), st.session_state.get("wycena_akpia_keys", set()))
    # sel/cab_sel/asix/budget/dev_budget policzone wyżej (sekcje 3, 7, 5, 9, 9a) -
    # PDF dostaje te same obiekty zamiast dobierać PLC/szafę/SCADA/kosztorys
    # jeszcze raz od zera.
    pdf_bio = create_pdf_report(
        devices, balance, project_label, platforma,
        sel, cab_sel, asix, budget, IO_TYPES,
        dev_budget=dev_budget,
    )

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button("📄 Raport (Word)", data=word_bio,
                           file_name=f"Raport_{project_label}.docx", width="stretch")
    with c2:
        st.download_button("📊 Zestawienie (Excel)", data=excel_bio,
                           file_name=f"Zestawienie_{project_label}.xlsx",
                           type="primary", width="stretch")
    with c3:
        st.download_button("📕 Raport (PDF)", data=pdf_bio,
                           file_name=f"Raport_{project_label}.pdf", width="stretch")
    with c4:
        if st.button("💾 Zapisz do historii", width="stretch",
                     help="Zapisuje bieżący stan (snapshot JSON + Word + Excel) do "
                          "historia_projektow/ i outputs/. Rób to świadomie, np. przed "
                          "wysłaniem oferty — nie każda zmiana suwaka musi zostać w archiwum."):
            try:
                save_outputs_to_disk(
                    project_label, devices, balance, word_bio, excel_bio,
                    hmi_entries=st.session_state.get("hmi_entries", []),
                    wycena_akpia_keys=st.session_state.get("wycena_akpia_keys", set()),
                    akpia_price_overrides=st.session_state.get("akpia_price_overrides", {}),
                )
                st.success("Zapisano do historii projektów.")
            except Exception as exc:
                st.error(f"Błąd zapisu plików: {exc}")


def main():
    st.set_page_config(page_title="Panel Inżyniera AKPiA", layout="wide")
    if not check_password():
        st.stop()

    for key, default in [("api_key_override", ""), ("model_id_override", ""), ("devices", None),
                         ("project_label", None), ("current_file", (None, None)),
                         ("extraction_diff", None),
                         ("devices_ai_alternative", None)]:
        if key not in st.session_state:
            st.session_state[key] = default

    # Przełączenie zakładki programowo (np. po "Wczytaj do bieżącej analizy"
    # w Historii) MUSI się zdarzyć PRZED instancjacją st.sidebar.radio(...,
    # key="tryb_dzialania") w render_sidebar() niżej - Streamlit zabrania
    # nadpisywania session_state klucza widgetu PO jego instancjacji w tym
    # samym przebiegu (StreamlitAPIException). Stąd pośredni klucz
    # "force_page", ustawiany przez wywołującego, zamiast bezpośrednio
    # tryb_dzialania - odkryte przez faktyczne uruchomienie appki.
    if st.session_state.get("force_page"):
        st.session_state.tryb_dzialania = st.session_state.force_page
        st.session_state.force_page = None

    page, settings = render_sidebar()
    api_key = get_api_key()

    st.title("Panel Inżyniera AKPiA")
    st.warning("**Human-in-the-Loop:** AI jedynie ekstrahuje listę urządzeń. "
               "Zliczanie I/O i dobór wykonuje deterministyczny rdzeń. "
               "Ostateczna weryfikacja należy do inżyniera.")

    if page == "Analiza Projektu":
        st.header("Analiza Projektu")
        st.caption("Wgraj zestawienie (Excel) i/lub dokumentację (PDF). "
                   "AI zbuduje listę urządzeń, rdzeń policzy sygnały I/O.")

        c1, c2 = st.columns(2)
        with c1:
            excel_file = st.file_uploader("Zestawienie urządzeń (Excel)", type=["xlsx", "xls"], key="xl")
        with c2:
            pdf_file = st.file_uploader("Dokumentacja (PDF)", type=["pdf"], key="pdf")

        # Wybór arkusza - plik może mieć kilka zakładek, aplikacja musi wiedzieć,
        # z której czytać (wcześniej zawsze brała pierwszą po cichu, co mylące
        # przy plikach wieloarkuszowych - patrz historia konwersacji).
        selected_sheet = 0
        if excel_file:
            sheets = list_excel_sheets(excel_file)
            if len(sheets) > 1:
                selected_sheet = st.selectbox(
                    "Arkusz do wczytania", sheets, index=0, key="sheet_select",
                    help="Plik ma więcej niż jeden arkusz. Wybierz właściwy - "
                         "domyślnie wczytywany jest pierwszy od lewej.",
                )
            elif sheets:
                selected_sheet = sheets[0]

        current_files = (excel_file.name if excel_file else None,
                         pdf_file.name if pdf_file else None, selected_sheet)
        # has_any_file: st.file_uploader NIE gwarantuje, że plik zostaje
        # "przypięty" po przełączeniu na inną stronę (Historia Projektów,
        # Ustawienia API) i powrocie - w praktyce wraca pusty. Bez tego
        # warunku sam powrót na tę stronę wyglądałby jak "usunięcie pliku"
        # i wyzerowałby już policzoną analizę, mimo że inżynier niczego
        # nie zmienił - zmierzone bezpośrednio: nawigacja Analiza -> Ustawienia
        # -> Analiza kasowała wyniki. Realny plik w uploaderze nadal
        # poprawnie wyzwala przeliczenie przy wgraniu innego pliku.
        # Ubocznie rozwiązuje też reload z historii (patrz "🔄 Wczytaj do
        # bieżącej analizy") - zaraz po nim w uploaderze też nic nie ma,
        # więc has_any_file=False i devices poprawnie zostają nietknięte.
        has_any_file = excel_file is not None or pdf_file is not None
        if has_any_file and current_files != st.session_state.current_file:
            st.session_state.devices = None
            st.session_state.current_file = current_files
            # Nowy plik -> poprzednie porównanie offline/AI dotyczyło innych
            # danych, więc traci sens i musi zniknąć razem z devices.
            st.session_state.extraction_diff = None
            st.session_state.devices_ai_alternative = None

        excel_df = None
        if excel_file:
            try:
                excel_df = read_excel_file(excel_file, sheet_name=selected_sheet)
                sheet_label = f"„{selected_sheet}”" if isinstance(selected_sheet, str) else ""
                total_rows = len(excel_df)
                if total_rows > 10:
                    st.markdown(f"**Podgląd arkusza {sheet_label}** — pokazano pierwsze 10 "
                               f"z **{total_rows}** wierszy. Analiza uwzględni wszystkie {total_rows}.")
                else:
                    st.markdown(f"**Podgląd arkusza {sheet_label}** — {total_rows} wierszy (wszystkie widoczne).")
                st.dataframe(excel_df.head(10), width="stretch")
            except ValueError as exc:
                st.error(str(exc))

        # --- Ścieżka bez AI: parsuj Excel bezpośrednio rdzeniem ---
        if excel_df is not None:
            if st.button("⚙ Policz I/O z Excela (bez AI)", width="stretch"):
                devices, warns = parse_devices(excel_df)
                st.session_state.devices = devices
                st.session_state.project_label = build_project_label(current_files[0], current_files[1], selected_sheet)
                st.session_state.extraction_diff = None  # nowa ścieżka - wyczyść poprzednie porównanie
                for w in warns:
                    st.warning(w)
                try:
                    persist_fresh_analysis(devices, st.session_state.project_label, settings)
                except Exception as exc:
                    st.error(f"Błąd zapisu plików: {exc}")

        # --- Ścieżka z AI: ekstrakcja JSON, potem rdzeń ---
        if not excel_file and not pdf_file:
            st.info("Wgraj co najmniej jeden plik, aby rozpocząć.")
        elif not api_key:
            st.error("Brak klucza API Gemini. Skonfiguruj w Ustawieniach lub secrets.toml.")
        elif st.button("🤖 Ekstrahuj urządzenia przez AI, potem policz I/O", type="primary",
                       width="stretch"):
            with st.spinner("AI ekstrahuje listę urządzeń..."):
                try:
                    records = run_extraction(
                        api_key=api_key, excel_df=excel_df,
                        pdf_bytes=pdf_file.getvalue() if pdf_file else None,
                        excel_filename=current_files[0], pdf_filename=current_files[1],
                    )
                    devices, warns = parse_ai_devices(records)
                    st.session_state.devices = devices
                    st.session_state.project_label = build_project_label(current_files[0], current_files[1], selected_sheet)
                    st.session_state.extraction_diff = None  # nowa ścieżka - wyczyść poprzednie porównanie
                    for w in warns:
                        st.warning(w)
                    try:
                        persist_fresh_analysis(devices, st.session_state.project_label, settings)
                    except Exception as save_exc:
                        st.error(f"Błąd zapisu plików: {save_exc}")
                except Exception as exc:
                    st.error(f"Błąd ekstrakcji: {exc}")

        # --- Ścieżka weryfikacyjna: obie metody naraz + porównanie (opcjonalna, kosztowa) ---
        if excel_df is not None:
            if not api_key:
                st.caption("⚖ Weryfikacja przez AI wymaga klucza API Gemini (patrz wyżej).")
            elif st.button("⚖ Policz + zweryfikuj przez AI", width="stretch",
                           help="Uruchamia OBIE metody (offline i AI) na tych samych danych "
                                "i pokazuje różnice. Zużywa dodatkowe zapytania do API - "
                                "użyj, gdy chcesz sprawdzić jakość ekstrakcji, nie przy każdej "
                                "iteracji."):
                with st.spinner("Liczenie offline + ekstrakcja AI + porównanie..."):
                    try:
                        devices_off, warns_off = parse_devices(excel_df)
                        records = run_extraction(
                            api_key=api_key, excel_df=excel_df,
                            pdf_bytes=pdf_file.getvalue() if pdf_file else None,
                            excel_filename=current_files[0], pdf_filename=current_files[1],
                        )
                        devices_ai, warns_ai = parse_ai_devices(records)

                        diff = compare_extractions(devices_off, devices_ai)
                        st.session_state.extraction_diff = diff

                        # Domyślnie do dalszej pracy bierzemy wynik OFFLINE (deterministyczny,
                        # bezpieczniejszy domyślny wybór) - inżynier może ręcznie przełączyć
                        # na wynik AI w sekcji porównania niżej.
                        st.session_state.devices = devices_off
                        st.session_state.project_label = build_project_label(
                            current_files[0], current_files[1], selected_sheet
                        )
                        st.session_state.devices_ai_alternative = devices_ai
                        for w in warns_off:
                            st.warning(w)
                        try:
                            persist_fresh_analysis(devices_off, st.session_state.project_label, settings)
                        except Exception as save_exc:
                            st.error(f"Błąd zapisu plików: {save_exc}")
                    except Exception as exc:
                        st.error(f"Błąd weryfikacji: {exc}")

        # --- Panel porównania (widoczny tylko po użyciu przycisku weryfikacji) ---
        if st.session_state.get("extraction_diff") is not None:
            render_extraction_diff_panel()

        # --- Wyniki (wspólne dla obu ścieżek) ---
        # Zapis do historii nie jest tu już automatyczny (patrz
        # persist_fresh_analysis) - świeża ekstrakcja zapisuje się raz przy
        # kliknięciu przycisku wyżej, a dalsze, świadome zapisy (np. po
        # dopracowaniu rabatów) robi przycisk "💾 Zapisz do historii"
        # w sekcji 11 wewnątrz render_results().
        if st.session_state.devices is not None:
            devices = st.session_state.devices
            balance = count_io(devices, reserve_percent=settings["reserve_percent"])
            render_results(devices, balance, st.session_state.project_label,
                          settings['platforma'], settings['rabaty'],
                          settings['cable_length'], settings['asix_factor'])

    elif page == "Historia Projektów":
        st.header("Archiwum (snapshoty JSON)")
        files = sorted([f for f in os.listdir(HISTORY_DIR) if f.endswith(".json")], reverse=True)
        if not files:
            st.info("Brak zapisanych projektów.")
        else:
            sel = st.selectbox("Wybierz analizę:", ["-- Wybierz --"] + files)
            if sel != "-- Wybierz --":
                with open(os.path.join(HISTORY_DIR, sel), encoding="utf-8") as f:
                    snap = json.load(f)
                st.json({"project": snap["project"], "reserve_percent": snap["reserve_percent"],
                         "balance_base": snap["balance_base"], "balance_reserved": snap["balance_reserved"]})
                st.dataframe(pd.DataFrame(snap["devices"]), width="stretch")

                st.caption(
                    "Dotąd historia była archiwum wyłącznie do odczytu — zmiana rabatów "
                    "czy rezerwy dla starego projektu wymagała ponownego wgrania "
                    "oryginalnego pliku. Przycisk niżej wczytuje urządzenia (razem z "
                    "wyceną AKPiA, HMI i ręcznymi cenami z tamtego zapisu, jeśli je miał) "
                    "z powrotem do bieżącej analizy — resztę (rabaty, rezerwę, platformę) "
                    "dobierasz od nowa w panelu bocznym."
                )
                if st.button("🔄 Wczytaj do bieżącej analizy", type="primary"):
                    st.session_state.devices = records_to_devices(snap["devices"])
                    st.session_state.project_label = snap["project"]
                    st.session_state.hmi_entries = snap.get("hmi_entries", [])
                    st.session_state.wycena_akpia_keys = set(snap.get("wycena_akpia_keys", []))
                    st.session_state.akpia_price_overrides = snap.get("akpia_price_overrides", {})
                    st.session_state.extraction_diff = None
                    st.session_state.devices_ai_alternative = None
                    st.session_state.force_page = "Analiza Projektu"
                    st.rerun()

    elif page == "Ustawienia API":
        st.header("Konfiguracja klucza Gemini")
        st.info("Klucz z `.streamlit/secrets.toml` ładuje się automatycznie. Tutaj nadpiszesz go na czas sesji.")
        st.session_state.api_key_override = st.text_input(
            "Tymczasowy klucz API", value=st.session_state.api_key_override, type="password")

        st.markdown("---")
        st.subheader("Model Gemini")
        st.caption(f"Domyślny model (w kodzie): `{GEMINI_MODEL}`. Zmiana obowiązuje "
                   "do końca tej sesji przeglądarki — bez edycji kodu i redeployu.")
        st.session_state.model_id_override = st.text_input(
            "Nadpisz model (opcjonalnie)", value=st.session_state.model_id_override,
            placeholder=GEMINI_MODEL,
            help="Zostaw puste, żeby użyć domyślnego modelu z kodu.",
        )
        st.caption(f"Aktualnie używany: `{get_model_id()}`")


if __name__ == "__main__":
    main()
